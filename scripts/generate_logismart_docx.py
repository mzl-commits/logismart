from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "documentacion_proyecto_logismart_actualizada.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_table(table, widths, header_fill="E8EEF5"):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row_index, row in enumerate(table.rows):
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(3)
                paragraph.paragraph_format.space_before = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(10.5)
            if row_index == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def set_page(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("Documentacion tecnica de LogiSmart")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4D, 0x78)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(
        "Estado actualizado del backend, despliegue, seguridad y operacion "
        "del sistema de clasificacion, almacenamiento y despacho asistido por AGV."
    )
    run.font.size = Pt(11)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def main():
    doc = Document()
    set_page(doc.sections[0])

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Heading 1"].font.name = "Calibri"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    styles["Heading 1"].paragraph_format.space_before = Pt(16)
    styles["Heading 1"].paragraph_format.space_after = Pt(8)
    styles["Heading 2"].font.name = "Calibri"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    styles["Heading 2"].paragraph_format.space_before = Pt(12)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)

    add_title(doc)

    doc.add_heading("1. Resumen ejecutivo", level=1)
    doc.add_paragraph(
        "LogiSmart es una plataforma para gestionar cajas, ubicaciones, "
        "despachos y movimientos de un AGV desde una interfaz web. Durante "
        "esta actualizacion se dejo la arquitectura alineada a produccion: "
        "PostgreSQL reemplazo a SQLite, Redis quedo habilitado para Channels, "
        "los servicios se ejecutan con Daphne bajo systemd y el sitio publico "
        "opera por HTTPS detras de Nginx."
    )
    doc.add_paragraph(
        "Tambien se corrigieron puntos criticos de seguridad y operacion: "
        "proteccion de la API interna, proteccion con X-API-Key para la API "
        "externa v1, rechazo de WebSockets anonimos, eliminacion de secretos "
        "embebidos, y una reserva atomica de ubicaciones para evitar dobles "
        "asignaciones cuando hay concurrencia."
    )

    doc.add_heading("2. Arquitectura del sistema", level=1)
    doc.add_paragraph(
        "El flujo principal del sistema conecta clientes web, API y hardware "
        "AGV mediante una arquitectura ASGI con componentes desacoplados."
    )

    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "Capa"
    table.rows[0].cells[1].text = "Componentes"
    table.rows[0].cells[2].text = "Responsabilidad"
    rows = [
        ("Acceso", "Nginx + HTTPS", "Termina TLS, redirige HTTP y sirve estaticos."),
        ("Aplicacion", "Daphne + Django + DRF + Channels", "UI, API, WebSocket y logica de negocio."),
        ("Datos", "PostgreSQL 15", "Persistencia principal del inventario y operaciones."),
        ("Tiempo real", "Redis 7", "Canales ASGI y mensajeria interna."),
        ("Integracion fisica", "Mosquitto MQTT + ESP32", "Comandos y telemetria del AGV."),
    ]
    for capa, componentes, responsabilidad in rows:
        cells = table.add_row().cells
        cells[0].text = capa
        cells[1].text = componentes
        cells[2].text = responsabilidad
    style_table(table, [Inches(1.15), Inches(2.05), Inches(3.3)])

    doc.add_heading("3. Mejoras tecnicas aplicadas", level=1)
    doc.add_heading("3.1 Seguridad", level=2)
    add_bullets(
        doc,
        [
            "DEBUG desactivado y SECRET_KEY obligatoria en produccion.",
            "ALLOWED_HOSTS y origenes CSRF configurados desde variables de entorno.",
            "Cookies seguras, redireccion HTTPS, HSTS y cabeceras defensivas activas.",
            "API interna protegida con SessionAuthentication + IsAuthenticated.",
            "API v1 protegida con X-API-Key usando comparacion en tiempo constante.",
            "WebSockets anonimos rechazados con cierre explicito.",
        ],
    )

    doc.add_heading("3.2 Disponibilidad y despliegue", level=2)
    add_bullets(
        doc,
        [
            "Migracion completa de produccion desde SQLite a PostgreSQL.",
            "Redis configurado como backend de Channels.",
            "Daphne y listener MQTT ejecutandose como servicios systemd.",
            "Nginx sirviendo HTTPS, proxy para ASGI y archivos estaticos.",
            "Permisos de estaticos corregidos para evitar 403 en produccion.",
        ],
    )

    doc.add_heading("3.3 Consistencia funcional", level=2)
    add_bullets(
        doc,
        [
            "Reserva atomica de ubicaciones con select_for_update para evitar carreras.",
            "Publicacion MQTT centralizada y con confirmacion de envio antes del cierre.",
            "Tests ajustados para entorno seguro con HTTPS y DEBUG=False.",
            "Suscripcion Stripe corregida eliminando expresiones invalidas del template.",
        ],
    )

    doc.add_heading("4. Flujo operativo", level=1)
    add_numbered(
        doc,
        [
            "Una caja entra al sistema en estado pendiente.",
            "El motor clasifica la caja y reserva la mejor ubicacion compatible.",
            "La ubicacion queda ocupada en la misma transaccion que cambia la caja a en_transito.",
            "Se calcula la ruta y se publica el comando al AGV por MQTT.",
            "Al confirmar la parada, la caja pasa a almacenada.",
            "Al despachar, la caja pasa a despachada y la ubicacion se libera.",
        ],
    )

    doc.add_heading("5. Variables y servicios criticos", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Elemento"
    table.rows[0].cells[1].text = "Uso"
    rows = [
        ("DJANGO_SECRET_KEY", "Clave base del proyecto Django."),
        ("DB_*", "Conexion a PostgreSQL."),
        ("REDIS_URL", "Backend de Channels."),
        ("MQTT_*", "Broker, credenciales y topics del AGV."),
        ("EXTERNAL_API_KEY", "Autorizacion para API externa v1."),
        ("STRIPE_*", "Checkout y webhook de suscripciones."),
        ("logismart.service", "Servicio principal ASGI."),
        ("logismart-mqtt.service", "Listener de telemetria MQTT."),
    ]
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
    style_table(table, [Inches(2.2), Inches(4.3)], header_fill="F2F4F7")

    doc.add_heading("6. Validacion realizada", level=1)
    add_bullets(
        doc,
        [
            "Pruebas locales: 28 de 28 OK.",
            "Pruebas en servidor con PostgreSQL y DEBUG=False: 28 de 28 OK.",
            "HTTP redirige correctamente a HTTPS.",
            "La raiz publica responde 302 hacia login cuando no hay sesion.",
            "La API /api/cajas/ responde 403 sin autenticacion.",
            "La API /api/v1/cajas responde 403 sin clave y 200 con X-API-Key valida.",
            "Los servicios logismart.service y logismart-mqtt.service quedaron activos.",
        ],
    )

    doc.add_heading("7. Riesgos y siguientes pasos", level=1)
    add_bullets(
        doc,
        [
            "Rotar las credenciales Stripe antes de activar cobros reales.",
            "Automatizar pg_dump con retencion y restauracion probada.",
            "Evaluar MQTT sobre TLS si el AGV operara fuera de una red confiable.",
            "Documentar versionado del firmware ESP32 y su contrato MQTT.",
        ],
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    set_page(doc.sections[-1])
    doc.add_heading("Anexo A. Rutas y archivos clave", level=1)
    add_bullets(
        doc,
        [
            "/home/yuri/proyecto_logistica",
            "/etc/systemd/system/logismart.service",
            "/etc/systemd/system/logismart-mqtt.service",
            "/etc/nginx/sites-available/logismart",
            "/var/www/logismart/static",
            "deploy/logismart.service",
            "deploy/logismart-mqtt.service",
            "deploy/nginx-logismart.conf",
        ],
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
